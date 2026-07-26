"""Turn a resume file into text.

Every reader here is behind a lazy import. v1 imported ``pypdf``, ``docx``,
``fitz`` and ``spacy`` at module scope, which made the whole package
un-importable unless ~2 GB of optional dependencies were present — the single
biggest obstacle to a ``uvx``-installable server. Import cost is paid only by
the caller who actually hands us a PDF.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from importlib.util import find_spec
from pathlib import Path

from careercraft.errors import DependencyMissing, ValidationFailed

SUPPORTED_SUFFIXES = frozenset({".pdf", ".docx", ".doc", ".txt", ".md", ".rtf"})

#: A PDF that yields less than this much text is probably a scan.
_OCR_TRIGGER_CHARS = 40


@dataclass(slots=True)
class LayoutLine:
    """One visual line of a PDF, with the font features used to spot headers."""

    text: str
    size: float
    bold: bool


@dataclass(slots=True)
class ExtractedDocument:
    text: str
    used_ocr: bool = False
    layout_lines: list[LayoutLine] = field(default_factory=list)
    source_name: str | None = None


def _has(module: str) -> bool:
    try:
        return find_spec(module) is not None
    except (ImportError, ValueError):  # pragma: no cover - malformed install
        return False


def extraction_backends() -> dict[str, bool]:
    """Which readers this install can use. Feeds the capabilities resource."""
    return {
        "pdf": _has("pypdf"),
        "pdf_layout": _has("fitz"),
        "docx": _has("docx"),
        "ocr": _has("pytesseract") and _has("pdf2image"),
    }


# ------------------------------------------------------------------ readers


def _read_pdf(path: Path) -> str:
    if not _has("pypdf"):
        raise DependencyMissing("Reading PDF resumes", "pdf")
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    if not _has("docx"):
        raise DependencyMissing("Reading Word resumes", "pdf", package="careercraft-mcp[pdf]")
    from docx import Document

    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _ocr_pdf(path: Path) -> str:
    from pdf2image import convert_from_path
    from pytesseract import image_to_string

    return "\n".join(image_to_string(page) for page in convert_from_path(str(path)))


def _pdf_layout_lines(path: Path) -> list[LayoutLine]:
    """Per-line text plus font size and boldness, via PyMuPDF.

    Many resumes mark sections purely visually — a 14pt bold "Experience" with
    no colon and no all-caps. Regex header detection misses those entirely, so
    this is what makes section splitting work on real documents.
    """
    if not _has("fitz"):
        return []
    import fitz

    lines: list[LayoutLine] = []
    doc = fitz.open(str(path))
    try:
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                for line in block.get("lines", []):
                    spans = line.get("spans", [])
                    text = "".join(s.get("text", "") for s in spans).strip()
                    if not text:
                        continue
                    sizes = [s.get("size", 0.0) for s in spans] or [0.0]
                    lines.append(
                        LayoutLine(
                            text=text,
                            size=sum(sizes) / len(sizes),
                            bold=any((s.get("flags", 0) & 1) == 1 for s in spans),
                        )
                    )
    finally:
        doc.close()
    return lines


# ------------------------------------------------------------- entry points


def extract_text(path: Path, *, enable_ocr: bool = True, fast: bool = False) -> ExtractedDocument:
    """Read ``path`` into text, with layout hints and OCR fallback for PDFs.

    ``fast=True`` skips both layout analysis and OCR. Section detection gets
    noticeably worse; use it only when throughput matters more than quality.
    """
    path = Path(path)
    if not path.is_file():
        raise ValidationFailed(
            f"{path} is not a readable file.",
            remedy="Pass an absolute path to a .pdf, .docx or .txt resume.",
        )

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValidationFailed(
            f"Cannot read {suffix or 'extension-less'} files.",
            remedy=f"Supported formats: {', '.join(sorted(SUPPORTED_SUFFIXES))}.",
        )

    if suffix == ".pdf":
        text = _read_pdf(path)
        used_ocr = False
        if not fast and enable_ocr and len(text.strip()) < _OCR_TRIGGER_CHARS:
            if extraction_backends()["ocr"]:
                try:
                    ocr_text = _ocr_pdf(path)
                except Exception:  # noqa: BLE001 - OCR is best-effort by design
                    ocr_text = ""
                if ocr_text.strip():
                    text, used_ocr = ocr_text, True
            elif not text.strip():
                raise DependencyMissing("Reading this scanned PDF", "ocr")
        layout = [] if fast else _pdf_layout_lines(path)
        return ExtractedDocument(text, used_ocr, layout, path.name)

    if suffix in (".docx", ".doc"):
        return ExtractedDocument(_read_docx(path), source_name=path.name)

    return ExtractedDocument(_read_txt(path), source_name=path.name)


def normalize_text(text: str) -> str:
    """Collapse the whitespace and bullet noise that PDF extraction leaves."""
    clean = text.replace("\r\n", "\n").replace("\r", "\n")
    clean = re.sub(r"[•‣◦⁃∙]", "•", clean)
    clean = re.sub(r"[ \t]+", " ", clean)
    return re.sub(r"\n{3,}", "\n\n", clean).strip()
